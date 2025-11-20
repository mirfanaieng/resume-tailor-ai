# src/formatter.py

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
import re

def parse_text_file(txt_path):
    """
    Parses a plain text resume into sections using headings.
    Example headings: Summary, Experience, Skills, Education
    Returns:
        dict: {section_name: content}
    """
    if not os.path.exists(txt_path):
        raise FileNotFoundError(f"File not found: {txt_path}")

    with open(txt_path, "r", encoding="utf-8") as f:
        text = f.read()

    # Regex pattern to detect section headings (case-insensitive)
    pattern = r"(Summary|Experience|Skills|Education):"
    splits = re.split(pattern, text, flags=re.IGNORECASE)
    
    sections = {}
    if splits[0].strip():  # any text before first heading
        sections["Introduction"] = splits[0].strip()

    for i in range(1, len(splits)-1, 2):
        heading = splits[i].strip().title()
        content = splits[i+1].strip()
        sections[heading] = content

    return sections

def create_formatted_resume_from_txt(txt_path, output_dir="data", filename=None):
    """
    Converts a tailored resume text file into a styled .docx file
    """
    sections = parse_text_file(txt_path)

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"tailored_resume_{timestamp}.docx"
    
    output_path = os.path.join(output_dir, filename)

    doc = Document()
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Add sections
    for section, content in sections.items():
        if not content.strip():
            continue  # skip empty sections

        doc.add_heading(section, level=1)
        para = doc.add_paragraph(content.strip())
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_after = Pt(6)

    doc.save(output_path)
    print(f"✅ Resume saved at: {output_path}")
    return output_path


if __name__ == "__main__":
    txt_path = "data/tailored_resume.txt"  # replace with your txt file
    create_formatted_resume_from_txt(txt_path)
